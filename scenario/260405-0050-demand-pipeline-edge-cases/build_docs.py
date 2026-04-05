"""Generate Elevate Foods Operations Playbook (DOCX) + Vendor Catalog Template (XLSX)."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

OUT = os.path.dirname(os.path.abspath(__file__))

# ── DOCX: Operations Playbook ──────────────────────────────────────────

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_heading("Elevate Foods — Operations Playbook", level=0)
doc.add_paragraph("Confidential — Internal Use Only")
doc.add_paragraph(f"Version 1.0 — April 2026")
doc.add_paragraph("")

# ── Section 1: Organization ──
doc.add_heading("1. Organization & Roles", level=1)

doc.add_heading("Team Structure", level=2)
doc.add_paragraph(
    "Elevate Foods operates with four pillars: Product (Tommy), Operations (System + Lawrence), "
    "Technology (Anik + AI Assistant), and Strategy (Owner). Each pillar operates autonomously "
    "with clear ownership boundaries."
)

roles = [
    ("Owner / Strategist", "You",
     "Strategic direction, systems design, growth planning, weekly oversight (approve cut orders, review metrics). "
     "NOT involved in: customer tickets, manual demand calculation, swap decisions, PO chasing."),
    ("Director of Product + Procurement", "Tommy Amorim",
     "Cheese selection, vendor relationships, PO approval, rotation calendar, quality oversight, supplier management. "
     "Reviews and approves auto-generated PO drafts every Monday. Plans cheese rotation 2-3 months ahead. "
     "Escalation point for product quality questions (taste, texture, sourcing)."),
    ("Developer", "Anik",
     "Builds and maintains AppyHour platform (Flask + pywebview). Implements features from specs. "
     "Currently building fulfillment tool — available for platform work after completion."),
    ("AI Assistant", "TBD",
     "Builds MCP tools, automation scripts, Gorgias reporting pipeline, Slack integrations. "
     "Works alongside Anik on tool-building and data analysis."),
    ("Customer Service", "Lawrence",
     "Full authority to handle CS tickets per Decision Authority Matrix. No owner approval needed for standard scenarios. "
     "Must tag every Gorgias ticket with Contact Reason + Resolution before closing. "
     "Escalates to Tommy for product quality, to Owner only for legal/PR/systemic issues."),
    ("Marketing", "Michelle",
     "Ad campaigns, seasonal box planning, SKU decisions. New process: post in #ads-team when scaling "
     "ad budget >2x with spend amount so inventory system can adjust safety buffer."),
    ("Fulfillment Center", "RMFG (Texas)",
     "Receives cut orders and POs. Cuts, portions, and packs subscription boxes. "
     "Needs accurate POs 3-5 days ahead of ship date. Ships Saturday (bulk) and Tuesday (first orders, reships)."),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Person"
hdr[2].text = "Responsibilities"
for role, person, resp in roles:
    row = table.add_row().cells
    row[0].text = role
    row[1].text = person
    row[2].text = resp

# ── Section 2: Weekly Cycle ──
doc.add_heading("2. Weekly Operations Cycle", level=1)

days = [
    ("Friday", [
        ("System", "Auto-pull inventory snapshot from Dropbox/RMFG"),
        ("System", "Recharge sync (queued charges, next 4 weeks)"),
        ("System", "Shopify sync (unfulfilled orders + rolling average)"),
        ("Owner", "Glance at weekend ship readiness — 5 min"),
    ]),
    ("Saturday", [
        ("RMFG", "Ships Saturday subscription orders"),
        ("System", "Auto-depletion from shipment XLSX → inventory journal entry"),
        ("System", "Tuesday projection calculated"),
    ]),
    ("Sunday", [
        ("System", "Weekly churn check (Recharge cancellations)"),
        ("System", "Growth multiplier recalculated from first-order trend"),
    ]),
    ("Monday", [
        ("System", "PO draft auto-generated (SKUs with runway < 2 weeks)"),
        ("System", "Slack notification to Tommy: 'X SKUs need POs — review in dashboard'"),
        ("Tommy", "Reviews and approves PO draft in dashboard — 15 min"),
        ("Owner", "Review shortage alerts, confirm PO approvals — 10 min"),
        ("System", "Approved POs emailed to suppliers"),
        ("Lawrence", "CS reship deadline: 5 PM ET for Saturday reships"),
    ]),
    ("Tuesday", [
        ("RMFG", "Ships Tuesday orders (first orders, reships, one-time boxes)"),
        ("System", "Auto-depletion from Tuesday shipment"),
        ("Lawrence", "CS reship deadline: 5 PM ET for next Tuesday reships"),
    ]),
    ("Wednesday", [
        ("System", "Cut order auto-generated: max(curation_floor, EWMA × growth_multiplier)"),
        ("System", "Swap suggestions generated if any SKU short"),
        ("Owner", "Review and approve cut order — 15 min"),
        ("Owner", "Approve swaps if needed — 15 min"),
        ("System", "Cut order XLSX exported and emailed to RMFG"),
        ("Tommy", "Monthly: update next month's curation recipes"),
    ]),
    ("Thursday", [
        ("Tommy", "Confirm RMFG received POs and cut order"),
        ("System", "Track PO acknowledgment status"),
    ]),
]

for day, tasks in days:
    doc.add_heading(day, level=2)
    for owner, task in tasks:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{owner}: ")
        run.bold = True
        p.add_run(task)

# ── Section 3: Demand Pipeline ──
doc.add_heading("3. How Demand Is Calculated", level=1)

doc.add_paragraph(
    "Demand comes from three sources, combined into a final number per SKU:"
)

doc.add_heading("Source 1: Recharge Queued Charges", level=2)
doc.add_paragraph(
    "Pulls all queued (not yet processed) subscription charges from Recharge API for the next 4 weeks. "
    "Each charge resolves to: box SKU → curation → PR-CJAM cheese + CEX-EC cheese + direct pickable SKUs. "
    "This is the primary demand signal for subscription boxes."
)

doc.add_heading("Source 2: Shopify Unfulfilled Orders", level=2)
doc.add_paragraph(
    "Pulls open/unfulfilled Shopify orders filtered by _SHIP_ tag. Adds direct SKU demand and "
    "prcjam/cexec counts on top of Recharge data. Also computes 8-week EWMA (exponential weighted "
    "moving average) from fulfilled order history for trend detection."
)

doc.add_heading("Source 3: MONG First-Order Projection", level=2)
doc.add_paragraph(
    "Counts 'Subscription First Order' tagged orders from last 3 days. Extrapolates daily rate to "
    "Monday midnight ET. Adds projected first orders to MONG curation prcjam/cexec counts."
)

doc.add_heading("Final Demand Formula", level=2)
doc.add_paragraph(
    "Final Demand per SKU = max(Curation Floor, EWMA Forecast × Growth Multiplier)"
)
doc.add_paragraph(
    "• Curation Floor = subscription_count × qty_per_box from active recipes. "
    "This guarantees you never order less than what recipes physically require.\n"
    "• EWMA Forecast = trend-aware weighted average of historical shipments. "
    "Recent weeks weighted 3x more than older weeks (alpha=0.3).\n"
    "• Growth Multiplier = week-over-week first-order acquisition trend (clamped ±20%)."
)

# ── Section 4: Subscription Base ──
doc.add_heading("4. Subscription Base (April 2026)", level=1)
doc.add_paragraph("Total active subscriptions: 13,489")

sub_data = [
    ("MDT", "2,027", "1,200", "3,227", "24%", "Set"),
    ("SPN", "1,127", "500", "1,627", "12%", "Set"),
    ("OWC", "974", "397", "1,371", "10%", "Set"),
    ("HHIGH", "881", "455", "1,336", "10%", "Set"),
    ("MS", "347", "474", "821", "6%", "Monthly"),
    ("ALPN", "465", "138", "603", "4.5%", "Set"),
    ("TRAY", "410", "—", "410", "3%", "Monthly"),
    ("ISUN", "254", "86", "340", "2.5%", "Set"),
    ("BYO", "145", "87", "232", "1.7%", "Set"),
    ("MONG", "136", "42", "178", "1.3%", "Set"),
    ("NMS", "28", "13", "41", "0.3%", "Monthly"),
    ("Legacy IDs", "~3,000", "—", "~3,000", "22%", "Unknown"),
]

table2 = doc.add_table(rows=1, cols=6)
table2.style = "Light Grid Accent 1"
hdr2 = table2.rows[0].cells
for i, h in enumerate(["Curation", "MED", "LGE", "Total", "% Base", "Type"]):
    hdr2[i].text = h
for row_data in sub_data:
    row = table2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph(
    "\nSet curations have stable recipes — demand is predictable via EWMA. "
    "Monthly curations (MS, NMS, TRAY, MED, LGE, CMED) get new recipes each month — "
    "demand needs auto-ramp buffer when new cheeses enter rotation. "
    "Legacy IDs (~3,000 subs on numeric Shopify variant IDs) need resolution to box types."
)

# ── Section 5: Tommy's Procurement Guide ──
doc.add_heading("5. Tommy's Procurement Guide", level=1)

doc.add_heading("Monday PO Review", level=2)
doc.add_paragraph(
    "Every Monday morning, the system generates a PO draft. Open the dashboard and review:\n"
    "• Each line shows: SKU, Vendor, Qty, Case Qty, Cost, Current Runway, Runway After PO\n"
    "• Click 'Approve' on lines that look right. Edit qty if you know something the system doesn't "
    "(e.g., supplier told you they're short, or you want to bump for a promo).\n"
    "• Approved POs auto-email to the vendor.\n"
    "• If a SKU shows red and no PO is possible (supplier can't deliver in time), flag it — "
    "the system will suggest a swap from surplus inventory."
)

doc.add_heading("Monthly Recipe Rotation", level=2)
doc.add_paragraph(
    "For monthly curations (MS, NMS, MED, LGE, CMED, TRAY):\n"
    "1. Open the Recipe Planner in the dashboard (Settings > Curation Recipes)\n"
    "2. Update next month's recipe — swap old cheeses for new ones\n"
    "3. System instantly shows demand impact: 'CH-NEW inherits CH-OLD demand (X units)'\n"
    "4. System auto-creates PO draft for new cheeses based on inherited demand\n"
    "5. Review and approve the auto-PO\n\n"
    "Plan 2-3 months ahead when possible. The earlier you enter recipes, the more lead time "
    "the system has to order from suppliers."
)

doc.add_heading("Vendor Catalog", level=2)
doc.add_paragraph(
    "Fill in the Vendor Catalog spreadsheet (separate XLSX file). This is critical — without it, "
    "the system can calculate HOW MUCH to order but not FROM WHOM or WHEN TO ORDER BY.\n\n"
    "Required fields per SKU: Vendor name, unit cost, case quantity (MOQ per order), "
    "minimum order quantity, lead time in days, shelf life in days, seasonal availability notes."
)

# ── Section 6: Lawrence's CS Guide ──
doc.add_heading("6. Lawrence's CS Operations", level=1)

doc.add_heading("Decision Authority (No Approval Needed)", level=2)
cs_auto = [
    ("Warm/spoiled arrival", "Full reship up to box value"),
    ("Missing item", "Add to next box OR partial reship"),
    ("Wrong box type", "Full reship"),
    ("Delayed >3 days", "Full reship (cold chain failure)"),
    ("Customer requests cancel", "Process cancellation"),
    ("Credit request", "$20 credit (standard resolution)"),
]
table3 = doc.add_table(rows=1, cols=2)
table3.style = "Light Grid Accent 1"
hdr3 = table3.rows[0].cells
hdr3[0].text = "Situation"
hdr3[1].text = "Lawrence's Authority"
for sit, auth in cs_auto:
    row = table3.add_row().cells
    row[0].text = sit
    row[1].text = auth

doc.add_heading("Must Escalate", level=2)
esc = [
    ("Product quality (taste/texture/sourcing)", "→ Tommy"),
    ("Allergen concern", "→ Owner (immediate)"),
    ("Legal threat", "→ Owner (immediate)"),
    ("Repeat customer (3+ tickets in 90 days)", "→ Owner"),
    ("Resolution >$30", "→ Owner"),
]
table4 = doc.add_table(rows=1, cols=2)
table4.style = "Light Grid Accent 1"
hdr4 = table4.rows[0].cells
hdr4[0].text = "Situation"
hdr4[1].text = "Escalation"
for sit, esc_to in esc:
    row = table4.add_row().cells
    row[0].text = sit
    row[1].text = esc_to

doc.add_heading("Gorgias Ticket Tagging (MANDATORY)", level=2)
doc.add_paragraph(
    "EVERY ticket must have Contact Reason + Resolution fields set BEFORE closing. No exceptions.\n\n"
    "Current tagging rate: 6%. Target: 90%+ within 2 weeks.\n\n"
    "Without tags, we cannot measure: CS cost per week, most common complaint types, "
    "which cheeses cause the most issues, carrier performance, or seasonal patterns. "
    "This data feeds Tommy's quality decisions and the Owner's strategic metrics."
)

# ── Section 7: Key Metrics ──
doc.add_heading("7. Key Metrics & Targets", level=1)

metrics = [
    ("Shortage incidents / week", "2-4", "0", "Count OOS/swap mentions in #core-team"),
    ("Shortage lead time", "0-3 days", "14+ days", "Days between alert and ship date"),
    ("Cut order time", "2-3 hours", "15 min", "Owner review + one-click approve"),
    ("PO turnaround", "Days of Slack DMs", "15 min Monday", "Tommy approves auto-draft"),
    ("Gorgias tag rate", "6%", "90%+", "Contact Reason + Resolution filled"),
    ("CS reships / week", "5-10", "<2", "Gorgias ticket count"),
    ("Forecast accuracy", "Unknown", "<15% MAE", "Forecast vs actual shipment"),
    ("New cheese stockout", "~1 per rotation", "0", "Shortages in first 2 weeks"),
    ("Subscriber growth", "~200/week", "Track trend", "First-order count weekly"),
]
table5 = doc.add_table(rows=1, cols=4)
table5.style = "Light Grid Accent 1"
hdr5 = table5.rows[0].cells
for i, h in enumerate(["Metric", "Current", "Target", "How to Measure"]):
    hdr5[i].text = h
for m in metrics:
    row = table5.add_row().cells
    for i, val in enumerate(m):
        row[i].text = val

# ── Section 8: Tray Launch ──
doc.add_heading("8. Tray Product — What Broke & What's Fixed", level=1)
doc.add_paragraph(
    "410 subscribers on AHB-MCUST-TRAY. When trays launched, inventory broke because:\n\n"
    "1. TRAY was not a recognized curation — demand pipeline returned None for tray boxes\n"
    "2. No PR-CJAM or CEX-EC demand counted for tray orders\n"
    "3. Depletion code explicitly SKIPPED items with '(tray)' in the product name\n"
    "4. Result: ~1,600-2,800 units of invisible demand AND invisible depletion\n\n"
    "Fixes applied (April 5, 2026):\n"
    "• Added TRAY to KNOWN_CURATIONS — resolve_curation_from_box_sku now returns 'TRAY'\n"
    "• Fixed depletion skip to only exclude packaging trays, not tray subscription products\n\n"
    "Still needed:\n"
    "• Define tray recipe in curation_recipes['TRAY'] (Tommy's task)\n"
    "• Configure PR-CJAM and CEX-EC for TRAY if applicable\n"
    "• Add TRAY to curation floor calculation"
)

# Save
docx_path = os.path.join(OUT, "Elevate_Foods_Operations_Playbook.docx")
doc.save(docx_path)
print(f"Saved: {docx_path}")

# ── XLSX: Vendor Catalog Template ──────────────────────────────────────

wb = openpyxl.Workbook()

# Style definitions
header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
alt_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
thin_border = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)

# ── Tab 1: Vendor Catalog ──
ws1 = wb.active
ws1.title = "Vendor Catalog"

headers1 = [
    "SKU", "Product Name", "Vendor", "Unit Cost ($)", "Case Qty",
    "MOQ (units)", "Lead Time (days)", "Shelf Life (days)",
    "Seasonal?", "Season Window", "Notes"
]
for col, h in enumerate(headers1, 1):
    cell = ws1.cell(row=1, column=col, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", wrap_text=True)
    cell.border = thin_border

# Sample rows for Tommy to fill in
sample_skus = [
    ("CH-EBRIE", "Époisses Brie", "", "", "", "", "14", "45", "No", "", ""),
    ("CH-SOT", "Sottocenere al Tartufo", "", "", "", "", "14", "60", "No", "", ""),
    ("CH-BLR", "Baked Lemon Ricotta", "", "", "", "", "7", "30", "No", "", ""),
    ("CH-MCPC", "McCall's Irish Porter", "", "", "", "", "21", "45", "No", "", ""),
    ("CH-CHALLER", "Challerhocker", "", "", "", "", "21", "60", "Yes", "Nov-Mar", "Very challenging natural rind"),
    ("CH-PVEC", "Piave Vecchio", "", "", "", "", "14", "90", "No", "", ""),
    ("CH-TOPR", "Toma Provence", "", "", "", "", "14", "45", "No", "", ""),
    ("MT-LONZ", "Lonza", "", "", "", "", "10", "90", "No", "", ""),
    ("MT-TUSC", "Tuscan Salami", "", "", "", "", "10", "90", "No", "", ""),
    ("AC-DTCH", "Dutch Stroopwafel", "", "", "", "", "14", "180", "No", "", ""),
    ("AC-DCRAN", "Dark Chocolate Cranberry", "", "", "", "", "14", "120", "No", "", ""),
    ("AC-BLBALS", "Balsamic Blueberry", "", "", "", "", "7", "60", "No", "", ""),
]
for r, row_data in enumerate(sample_skus, 2):
    for c, val in enumerate(row_data, 1):
        cell = ws1.cell(row=r, column=c, value=val)
        cell.border = thin_border
        if r % 2 == 0:
            cell.fill = alt_fill

# Set column widths
widths1 = [12, 25, 20, 12, 10, 12, 14, 14, 10, 14, 30]
for i, w in enumerate(widths1, 1):
    ws1.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

# ── Tab 2: Supplier Directory ──
ws2 = wb.create_sheet("Supplier Directory")

headers2 = [
    "Vendor Name", "Contact Name", "Email", "Phone",
    "Products Supplied", "Payment Terms", "Min Order ($)",
    "Shipping Method", "Lead Time Range", "Notes"
]
for col, h in enumerate(headers2, 1):
    cell = ws2.cell(row=1, column=col, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", wrap_text=True)
    cell.border = thin_border

# Known suppliers from Slack
suppliers = [
    ("Forever Cheese", "", "", "", "European imports (Époisses, Piave, etc.)", "", "", "", "14-21 days", ""),
    ("WBC", "", "", "", "Specialty cheeses (Challerhocker, etc.)", "", "", "", "14-21 days", ""),
    ("VT Salumi", "", "", "", "Charcuterie (Lonza, etc.)", "", "", "", "7-14 days", ""),
    ("Linscott", "", "", "", "", "", "", "", "", ""),
    ("RMFG", "", "", "", "Fulfillment center — cutting, portioning, packing", "", "", "", "3-5 days", "Texas-based"),
]
for r, row_data in enumerate(suppliers, 2):
    for c, val in enumerate(row_data, 1):
        cell = ws2.cell(row=r, column=c, value=val)
        cell.border = thin_border
        if r % 2 == 0:
            cell.fill = alt_fill

widths2 = [18, 16, 25, 15, 30, 15, 14, 16, 16, 30]
for i, w in enumerate(widths2, 1):
    ws2.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

# ── Tab 3: Rotation Calendar ──
ws3 = wb.create_sheet("Rotation Calendar")

headers3 = [
    "Month", "Curation", "Slot", "Old SKU", "New SKU",
    "Vendor", "PO Qty Needed", "Lead Time", "Order By Date", "Status"
]
for col, h in enumerate(headers3, 1):
    cell = ws3.cell(row=1, column=col, value=h)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", wrap_text=True)
    cell.border = thin_border

# Example row
example = ("May 2026", "MDT", "Cheese 1", "CH-SOT", "CH-NEW", "TBD", "3,227", "14 days", "Apr 17", "Planning")
for c, val in enumerate(example, 1):
    cell = ws3.cell(row=2, column=c, value=val)
    cell.border = thin_border

widths3 = [12, 12, 10, 14, 14, 14, 14, 12, 14, 12]
for i, w in enumerate(widths3, 1):
    ws3.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

# Save
xlsx_path = os.path.join(OUT, "Elevate_Foods_Vendor_Catalog.xlsx")
wb.save(xlsx_path)
print(f"Saved: {xlsx_path}")
