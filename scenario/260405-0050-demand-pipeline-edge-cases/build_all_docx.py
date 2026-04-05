"""Convert all markdown plans into DOCX files."""
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))


def md_to_docx(md_path, docx_path, title=None):
    """Convert a markdown file to a styled DOCX."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_table = False
    table_rows = []
    in_code = False
    code_lines = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        # Parse markdown table
        headers = [c.strip() for c in table_rows[0].strip("|").split("|")]
        data = []
        for row in table_rows[2:]:  # skip separator
            cells = [c.strip() for c in row.strip("|").split("|")]
            data.append(cells)

        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for row_data in data:
            row = t.add_row().cells
            for i, val in enumerate(row_data):
                if i < len(row):
                    row[i].text = val

        table_rows = []
        in_table = False

    def flush_code():
        nonlocal code_lines, in_code
        if code_lines:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_lines = []
        in_code = False

    for line in lines:
        stripped = line.rstrip("\n")

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(stripped)
            continue

        # Tables
        if "|" in stripped and stripped.strip().startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table()

        # Headings
        if stripped.startswith("# "):
            if title and stripped.lstrip("# ").strip() == title:
                doc.add_heading(title, level=0)
            else:
                doc.add_heading(stripped.lstrip("# ").strip(), level=0)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped.lstrip("## ").strip(), level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped.lstrip("### ").strip(), level=2)
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped.lstrip("#### ").strip(), level=3)
            continue

        # Bold line (standalone)
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold within bullets
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part.strip("*"))
                    run.bold = True
                else:
                    p.add_run(part)
            continue

        # Numbered lists
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            text = m.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*[^*]+\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part.strip("*"))
                    run.bold = True
                else:
                    p.add_run(part)
            continue

        # Checkbox items
        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x] ")
            text = stripped[6:].strip()
            prefix = "[x] " if checked else "[ ] "
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(prefix + text)
            continue

        # Empty lines
        if not stripped.strip():
            continue

        # Regular paragraph with inline bold
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part.strip("*"))
                run.bold = True
            else:
                p.add_run(part)

    # Flush remaining
    if in_table:
        flush_table()
    if in_code:
        flush_code()

    doc.save(docx_path)
    print(f"  Created: {os.path.basename(docx_path)}")


# Convert all plans
plans = [
    ("business-synthesis.md", "Business_Synthesis.docx", "Elevate Foods — Business Synthesis & Autopilot Roadmap"),
    ("operations-plan.md", "Operations_Plan.docx", "Elevate Foods — Operations Plan & Software Roadmap"),
    ("ops-reorg.md", "Ops_Reorg.docx", "Elevate Foods — Operations Reorganization"),
    ("autopilot-plan.md", "Autopilot_Plan.docx", "Autopilot Forecasting Plan"),
    ("summary.md", "Scenario_Analysis_Summary.docx", "Demand Pipeline Scenario Exploration — Summary"),
]

print("Converting markdown plans to DOCX...\n")
for md_name, docx_name, title in plans:
    md_path = os.path.join(OUT, md_name)
    docx_path = os.path.join(OUT, docx_name)
    if os.path.exists(md_path):
        md_to_docx(md_path, docx_path, title)
    else:
        print(f"  SKIP: {md_name} not found")

print(f"\nDone. {len(plans)} documents in {OUT}")
