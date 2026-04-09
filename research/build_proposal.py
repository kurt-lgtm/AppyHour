"""Generate gift card migration proposal as .docx for CEO review."""

from __future__ import annotations

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# -- Title --
title = doc.add_heading("Gift Card Order Editability", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("Proposal: Migrate to Discount-Based Gifting")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph("Prepared by: Operations Team")
doc.add_paragraph("Date: April 2026")
doc.add_paragraph("")

# -- Executive Summary --
doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "Our gift subscription program generates recurring revenue and customer acquisition. "
    "However, a technical limitation in Shopify prevents us from editing any order that was "
    "paid with a gift card. This means we cannot add bonus items, swap products for dietary "
    "restrictions, or make any modifications to gift redemption orders before they ship."
)
doc.add_paragraph(
    "This proposal outlines a validated solution that preserves the gift subscription "
    "experience for customers while giving our operations team full control over order contents. "
    "The change is invisible to customers and requires no changes to our pricing or product offerings."
)

# -- The Solution --
doc.add_heading("The Solution", level=1)
doc.add_paragraph(
    "Replace the gift card payment mechanism with Recharge discount codes. "
    "Instead of issuing a Shopify gift card when someone buys a gift subscription, "
    "we create a discount code that covers the recipient's charges for the gift period."
)

# comparison table
doc.add_heading("How It Works: Before vs. After", level=2)
table = doc.add_table(rows=5, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["", "Current (Gift Card)", "Proposed (Discount Code)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

rows_data = [
    ["Customer buys gift", "Shopify issues a gift card", "System creates a discount code"],
    ["Recipient's monthly charge", "Gift card balance deducted", "Discount reduces the charge amount"],
    ["Payment method on order", "Gift card (locks the order)", "Credit card at reduced price ($0 if fully covered)"],
    ["Can we edit the order?", "NO", "YES"],
]
for r, row_data in enumerate(rows_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph("")
doc.add_paragraph(
    "The customer experience is identical — gift buyers purchase the same product at the same price, "
    "and recipients see their charges covered each month. The only difference is behind the scenes: "
    "orders are paid by the recipient's credit card (at $0 or a reduced amount) instead of by a gift card, "
    "which means we can edit them freely."
)

# -- Validation --
doc.add_heading("Validation", level=1)
doc.add_paragraph(
    "Every aspect of this solution has been verified against the official Shopify and Recharge documentation:"
)

val_table = doc.add_table(rows=6, cols=2)
val_table.style = "Light Grid Accent 1"
val_table.alignment = WD_TABLE_ALIGNMENT.CENTER
val_headers = ["Claim", "Verified?"]
for i, h in enumerate(val_headers):
    cell = val_table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

val_rows = [
    ["Orders with discount codes are editable in Shopify", "YES — confirmed via Shopify GraphQL API docs"],
    [
        "Discounts can be applied automatically without customer action",
        "YES — Recharge API supports backend-only application",
    ],
    ["Discounts can last for multiple months (e.g., 3 charges)", "YES — Recharge API field controls charge count"],
    ["$0 orders are created and remain editable", "YES — valid payment method required but $0 works"],
    [
        "The entire process can run inside Shopify + Recharge (no external servers)",
        "YES — using Mechanic automation app",
    ],
]
for r, row_data in enumerate(val_rows):
    for c, val in enumerate(row_data):
        val_table.rows[r + 1].cells[c].text = val

# -- Implementation Plan --
doc.add_heading("Implementation Plan", level=1)

steps = [
    (
        "Step 1: Proof of Concept (Day 1)",
        [
            "Create a new test customer in Recharge with a test subscription (do NOT use a real customer)",
            "Manually create a test discount code in Recharge (e.g., 'TEST-GIFT-001', fixed amount, 1 charge)",
            "Apply the discount to the test subscription",
            "Trigger a charge for the test subscription",
            "Verify the resulting Shopify order is editable (Edit button visible in Shopify Admin)",
            "Add a line item to the order to confirm full edit capability",
            "This single test validates the entire approach — if it works, everything else is execution",
            "Clean up: cancel the test subscription and disable the test discount after verification",
        ],
    ),
    (
        "Step 2: Disable Gift Card Auto-Apply (Day 1)",
        [
            "Turn off the Recharge setting that automatically applies gift card balances to recurring charges",
            "This immediately stops new orders from being locked",
            "Existing gift card balances remain on customer accounts (migrated in Step 3)",
        ],
    ),
    (
        "Step 3: Migrate Existing Gift Card Balances (Day 1-2)",
        [
            "Identify all customers with remaining gift card balances",
            "For each customer: create an equivalent discount code in Recharge",
            "Apply the discount to their subscription",
            "Disable the old Shopify gift card",
            "Estimated volume: small number of active gift cards to migrate",
        ],
    ),
    (
        "Step 4: Create New Gift Product (Day 2)",
        [
            "Create a new Shopify product for gift subscriptions (same name, same price)",
            "The only change: product type is not 'Gift Card' — this prevents the payment lock",
            "Archive the old gift card product so no new purchases use the old flow",
            "Update any website links to point to the new product",
        ],
    ),
    (
        "Step 5: Set Up Automation (Day 2-3)",
        [
            "Configure Mechanic (our existing Shopify automation app) to handle new gift purchases automatically",
            "When a gift order comes in: Mechanic creates a discount code and applies it to the recipient",
            "Daily check: if the recipient hasn't signed up yet, Mechanic retries the next day",
            "Weekly check: when a discount is fully used, Mechanic applies any queued gifts",
            "All automation runs inside Shopify — no external servers or hosting needed",
        ],
    ),
    (
        "Step 6: Verify (Day 7-14)",
        [
            "After the next billing cycle, confirm all migrated orders are editable",
            "Verify discounts are applying correctly and counting down properly",
            "Confirm the automation handles new gift purchases end-to-end",
        ],
    ),
]

for title, bullets in steps:
    doc.add_heading(title, level=2)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

# -- Customer Impact --
doc.add_heading("Customer Impact", level=1)
doc.add_paragraph("This change is invisible to customers. The gift buying and receiving experience remains identical:")
impacts = [
    "Gift buyers: Same product, same price, same checkout flow",
    "Gift recipients: Same experience — they see their monthly charge covered by the gift",
    "No emails, notifications, or communication needed",
    "No changes to pricing, packaging, or product offerings",
]
for b in impacts:
    doc.add_paragraph(b, style="List Bullet")

# -- Risk & Rollback --
doc.add_heading("Risk Assessment", level=1)

doc.add_heading("Low Risk", level=2)
doc.add_paragraph("The old and new systems can run side-by-side. If anything goes wrong:")
rollback = [
    "Turn gift card auto-apply back on in Recharge (5 minutes)",
    "Re-enable the disabled gift cards",
    "Switch back to the old gift card product",
    "Discount codes continue working in parallel — nothing breaks",
]
for b in rollback:
    doc.add_paragraph(b, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph("This is not a one-way door. Both systems coexist, and we can revert at any time.").runs[
    0
].bold = True

doc.add_heading("Known Limitation", level=2)
doc.add_paragraph(
    "Recharge allows only one discount code per subscription at a time. This means a gift "
    "customer cannot stack a gift discount with a promotional discount simultaneously. "
    "However, since orders are now editable, promotional adjustments can be applied directly "
    "to the order after creation — the same way we handle bonus items today."
)

# -- Timeline & Cost --
doc.add_heading("Timeline & Cost", level=1)

tl_table = doc.add_table(rows=7, cols=3)
tl_table.style = "Light Grid Accent 1"
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tl_headers = ["What", "When", "Effort"]
for i, h in enumerate(tl_headers):
    cell = tl_table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

tl_rows = [
    ["Proof of concept test", "Day 1", "30 minutes"],
    ["Disable gift card auto-apply", "Day 1", "5 minutes"],
    ["Migrate existing balances", "Day 1-2", "1-2 hours"],
    ["Create new gift product", "Day 2", "30 minutes"],
    ["Set up Mechanic automation", "Day 2-3", "2-3 hours"],
    ["Verify after billing cycle", "Day 7-14", "30 minutes"],
]
for r, row_data in enumerate(tl_rows):
    for c, val in enumerate(row_data):
        tl_table.rows[r + 1].cells[c].text = val

doc.add_paragraph("")
p = doc.add_paragraph("Total hands-on work: approximately ")
run = p.add_run("one day")
run.bold = True
p.add_run(
    ". Then one billing cycle to verify. No additional software costs — "
    "uses existing Shopify, Recharge, and Mechanic subscriptions."
)

# -- Recommendation --
doc.add_heading("Recommendation", level=1)
doc.add_paragraph(
    "We recommend proceeding with this migration. The proof of concept (Step 1) can be "
    "completed in 30 minutes and will definitively confirm the solution works. The full "
    "migration takes approximately one day of hands-on work, has a clean rollback plan, "
    "and is invisible to customers."
)
doc.add_paragraph(
    "The result: every gift subscription order becomes fully editable, allowing us to "
    "maintain our standard quality and curation processes across all orders — gift and non-gift alike."
)

# -- Save --
output_path = os.path.join(os.path.dirname(__file__), "Gift_Card_Migration_Proposal.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
