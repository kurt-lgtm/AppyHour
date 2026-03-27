"""Create cheesemonger onboarding document and upload to Google Drive."""
import os
import sys

# Use the project's google_integration module
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "GelPackCalculator"))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x2E, 0x4A, 0x62)

# ── Title ───────────────────────────────────────────────────────────
doc.add_heading("Elevate Foods / Appy Hour — Procurement & Product Guide", level=0)
p = doc.add_paragraph("Prepared for: New Cheesemonger — Procurement Team")
p.runs[0].bold = True
doc.add_paragraph(
    "This document covers everything you need to know about our product line, "
    "SKU system, weekly fulfillment cycle, inventory management, and procurement "
    "decision-making. It is your go-to reference for understanding what we ship, "
    "how we organize it, and what drives our weekly production orders."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Business Overview", level=1)
doc.add_paragraph(
    "Elevate Foods operates the Appy Hour subscription box — a curated monthly "
    "box of artisan cheeses, charcuterie meats, crackers, and accompaniments "
    "shipped directly to consumers nationwide. Every box ships cold (insulated "
    "box + gel packs) and contains a mix of cheese, meat, and accessory items "
    "selected by our curation team."
)
doc.add_paragraph(
    "We fulfill from RMFG in Texas (primary fulfillment center) with a secondary "
    "warehouse in Woburn, MA for receiving bulk product, processing, and "
    "cross-docking to Texas."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. SKU System & Taxonomy", level=1)
doc.add_paragraph(
    "Every product in our system has a SKU (Stock Keeping Unit) with a prefix "
    "that tells you exactly what category it belongs to. Understanding these "
    "prefixes is essential."
)

doc.add_heading("SKU Prefix Reference", level=2)
t = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text = "Prefix"
hdr[1].text = "Category"
hdr[2].text = "Examples"
hdr[3].text = "Notes"
for row_data in [
    ("CH-", "Cheese", "CH-BLR, CH-MAFT, CH-EBRIE, CH-FOWC, CH-BRZ",
     "Core cheese products. Sliced, wrapped, and labeled. Sourced as wheels and cut in-house or purchased pre-cut."),
    ("MT-", "Meat / Charcuterie", "MT-SOP, MT-LONZ, MT-TUSC, MT-BRAS",
     "Cured meats. Typically pre-sliced from vendors."),
    ("AC-", "Accessories", "AC-BLBALS, AC-TCRISP, AC-DTCH, AC-PRPE, AC-SMAL",
     "Crackers, spreads, nuts, dried fruit, honey, mustard, etc. Some processed from bulk at Woburn."),
    ("AHB-", "Box Product", "AHB-MCUST-MONG, AHB-LGE, AHB-MED",
     "The subscription box itself. Not a physical SKU you procure — it's the container/subscription record."),
    ("PR-CJAM", "Cheese & Jam Pairing", "PR-CJAM-MONG, PR-CJAM-GEN",
     "Bonus pairing item (1 per box). Each curation maps to a unique cheese. Included free."),
    ("CEX-EC", "Extra Cheese Add-on", "CEX-EC-MDT, CEX-EC-OWC",
     "~40% of subscribers add an extra cheese. Resolved per-curation. Sometimes splits across 2 cheeses."),
    ("EX-EA / EX-EM", "Paid Extras", "EX-EA, EX-EM",
     "Paid extra accessory or extra meat. Customer-purchased add-ons."),
    ("BL-", "Paid Bundle", "BL-BLR4",
     "Paid product bundles (e.g., BL-BLR4 = 4x CH-BLR + 2x AC-BLBALS). Expands to multiple food items."),
    ("PK-", "Tasting Guide", "PK-GUIDE",
     "Printed tasting guide insert. Not a food item."),
]:
    row = t.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading("What Counts as a 'Food Item'", level=2)
doc.add_paragraph(
    "Only CH-, MT-, and AC- prefixed SKUs count toward the food item total in a box. "
    "Everything else (AHB-, PR-CJAM, CEX-EC, PK-, BL-, EX-) is either the box itself, "
    "a pairing, an add-on, or packaging — not counted toward the food total."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. Box Types & What Goes In Them", level=1)

doc.add_heading("Custom Curated Boxes (your primary focus)", level=2)
doc.add_paragraph(
    "These are the core subscription product. Food items are selected monthly by "
    "our curation team using a dedicated tool."
)
t2 = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0].cells
hdr2[0].text = "SKU Pattern"
hdr2[1].text = "Size"
hdr2[2].text = "Food Items"
hdr2[3].text = "Composition"
for row_data in [
    ("AHB-MCUST-{curation}", "Medium", "7 items",
     "2 cheeses + 2 meats + crackers + 2 accessories"),
    ("AHB-LCUST-{curation}", "Large", "9 items",
     "3 cheeses + 3 meats + crackers + 2 accessories"),
]:
    row = t2.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_paragraph(
    "The curation suffix (e.g., MONG, MDT, OWC, SPN, ALPT, ISUN, HHIGH) tells "
    "you which flavor profile / recipe the box follows. Each curation has a fixed "
    "recipe of 7 or 9 food items that rotates monthly."
)

doc.add_heading("Monthly Boxes", level=2)
doc.add_paragraph(
    "Non-custom boxes where food items are assigned once at the start of the month "
    "(not by the curation tool). They follow the same item counts as custom boxes."
)
bullets = [
    "AHB-MED — Medium monthly box (7 food items)",
    "AHB-LGE — Large monthly box (9 food items)",
    "AHB-CMED — Cheese-only medium box (no meats)",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Promo Boxes", level=2)
doc.add_paragraph(
    "Promotional subscription products (e.g., 'AppyHour Box + FREE Artisan Cheese "
    "& Jam Pairings for Life'). These look different in our system — the Shopify "
    "line item often has a blank SKU — but they ship the same food as monthly boxes. "
    "The bundle should contain AHB-MED or AHB-LGE + PR-CJAM-GEN."
)

doc.add_heading("Free Brie for a Year", level=2)
doc.add_paragraph(
    "A promo variant where the customer gets CH-EBRIE (Échiré Brie) instead of "
    "the standard PR-CJAM cheese pairing. If a box has CH-EBRIE, we do NOT add "
    "PR-CJAM-GEN — they are mutually exclusive."
)

doc.add_heading("Specialty Boxes", level=2)
doc.add_paragraph(
    "AHB-X{suffix} (e.g., AHB-XSPN, AHB-XMONG) are one-time specialty/gift boxes. "
    "These follow their own recipes and are excluded from standard error checks."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Curations — The Flavor Profiles", level=1)
doc.add_paragraph(
    "Each curation is a themed flavor profile with a fixed recipe. The curation "
    "determines which cheeses, meats, and accessories go in a box."
)
doc.add_heading("Active Curations", level=2)
curations = [
    ("MONG", "Mongolian-inspired / bold flavors"),
    ("MDT", "Mediterranean"),
    ("OWC", "Old World Classic"),
    ("SPN", "Spanish-inspired"),
    ("ALPT / ALPN", "Alpine"),
    ("ISUN", "Italian Sun"),
    ("HHIGH", "Highland / hearty"),
    ("NMS", "Newer curation track (~167 subscribers)"),
    ("BYO", "Build Your Own (customer-selected items, ~28 subscribers)"),
    ("SS", "Separate curation track (~27 subscribers)"),
]
for code, desc in curations:
    doc.add_paragraph(f"{code} — {desc}", style="List Bullet")

doc.add_heading("Curation Rotation", level=2)
doc.add_paragraph(
    "Curations rotate monthly. Each month, a curation gets a new recipe. The "
    "curation suffix on the box SKU is always the LAST segment after splitting "
    "on dashes. For example: AHB-MCUST-CORS-MDT → curation is MDT."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Cheese Procurement — What You Need to Know", level=1)

doc.add_heading("Sourcing: Wheels vs Pre-Cut", level=2)
doc.add_paragraph(
    "Most of our cheeses arrive as whole wheels. We track wheel inventory separately "
    "from sliced inventory. The conversion factor is:"
)
doc.add_paragraph(
    "Potential slices = wheel weight (lbs) × wheel count × 2.67",
    style="List Bullet"
)
doc.add_paragraph(
    "For example, a 10 lb wheel yields approximately 26-27 sliced portions. When "
    "we need more sliced inventory, we cut wheels — the Wednesday cut order tells "
    "you exactly how many wheels to cut per SKU."
)

doc.add_heading("Wheel Inventory Tracking", level=2)
doc.add_paragraph(
    "Wheel inventory is tracked in our system with: weight per wheel (lbs), "
    "wheel count, and target sliced SKU. The system calculates potential supply "
    "from uncut wheels and factors it into reorder decisions."
)

doc.add_heading("PR-CJAM Cheese Assignments", level=2)
doc.add_paragraph(
    "Each curation maps to exactly one cheese for the PR-CJAM (bonus pairing). "
    "This cheese must be UNIQUE across curations — no two curations can share the "
    "same PR-CJAM cheese. When making procurement decisions, remember that PR-CJAM "
    "demand adds 1 unit per box on top of recipe demand."
)

doc.add_heading("CEX-EC Extra Cheese (~40% of Large boxes)", level=2)
doc.add_paragraph(
    "About 40% of subscribers (mainly large box) have the Curator's Extra Cheese "
    "add-on. This adds one cheese per box. Assignments follow an adjacency rule:"
)
doc.add_paragraph(
    "Curations are ordered: MONG, MDT, OWC, SPN, ALPT, ISUN, HHIGH. "
    "The CEX-EC cheese for a curation must NOT overlap with recipe cheeses "
    "from curations within 2 positions in this list. This prevents flavor repetition.",
)
doc.add_paragraph(
    "Some curations split CEX-EC across two cheeses by percentage (e.g., MDT: "
    "64% MCPC + 36% MSMG). Factor both into demand forecasting."
)

doc.add_heading("Substitution Rules", level=2)
doc.add_paragraph(
    "When a cheese is short, we have established substitution families:"
)
subs = [
    ("Brie family", "CH-TTBRIE, CH-TIP, CH-EBRIE are interchangeable bries"),
    ("Porter family", "CH-MCPC and CH-IPRW are both Irish porter cheeses"),
    ("Alpine/semi-hard", "CH-BARI (Barista) can substitute for CH-ALPHA (Alpha Tolman)"),
    ("Cheddar variants", "Check recipe context before substituting — flavor profiles differ"),
]
for name, desc in subs:
    doc.add_paragraph(f"{name}: {desc}", style="List Bullet")

doc.add_paragraph(
    "\nBefore assigning a substitute cheese, always check: Is it already a CEX-EC "
    "or PR-CJAM for any curation? If so, does the combined demand (substitution + "
    "assignment) still fit within available surplus?"
)

doc.add_heading("Excluded SKUs", level=2)
doc.add_paragraph(
    "CH-MAFT is NEVER assigned as a PR-CJAM or CEX-EC cheese — it is on the "
    "permanent exclusion list. If you see CH-MAFT on an order outside its native "
    "curation recipe, that's an error (likely stale data from a prior month)."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. Weekly Fulfillment Cycle", level=1)

doc.add_heading("Schedule", level=2)
t3 = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = t3.rows[0].cells
hdr3[0].text = "Day"
hdr3[1].text = "Action"
hdr3[2].text = "Your Role"
for row_data in [
    ("Friday", "Dropbox inventory snapshot received (planning numbers)",
     "Review inventory levels, flag shortages"),
    ("Saturday", "Fulfillment #1 — main shipment (largest batch)",
     "Ensure all cheese is cut and available"),
    ("Monday", "Updated inventory snapshot received",
     "Reconcile post-Saturday numbers"),
    ("Tuesday", "Fulfillment #2 — smaller batch",
     "Cover any remaining orders"),
    ("Wednesday AM", "Submit production/cut order to RMFG",
     "THIS IS YOUR KEY DELIVERABLE — the weekly cut order"),
]:
    row = t3.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading("Wednesday Cut Order — Your Weekly Deliverable", level=2)
doc.add_paragraph(
    "Every Wednesday morning, you generate the cut order that tells RMFG which "
    "wheels to cut and how many slices we need. This is driven by:"
)
sources = [
    "Recharge queued charges — subscription renewals scheduled for the upcoming week",
    "Shopify orders — one-time purchases and first orders tagged with the ship week",
    "First-order projection — we multiply active-curation first orders by 3x to account "
    "for orders that trickle in throughout the week",
    "Current sliced inventory — what's already on hand",
    "Wheel inventory — potential supply from uncut wheels",
]
for s in sources:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("Cut Order Math", level=2)
doc.add_paragraph(
    "For each cheese SKU:\n"
    "  Total Demand = Recharge demand + Shopify demand\n"
    "  Net Position = Sliced On Hand + Wheel Potential − Total Demand\n"
    "  If Net Position < 0: Wheels to Cut = ceil(shortage / (wheel_weight_lbs × 2.67))\n"
    "  Cap at available wheel count"
)

doc.add_heading("Cross-Dock Timeline (Woburn → Texas)", level=2)
doc.add_paragraph(
    "If product needs to come from Woburn to the Texas fulfillment center:"
)
cross = [
    "Thursday: Bulk arrives at Woburn",
    "Friday: Cross-dock pickup (Woburn → RMFG TX)",
    "+9 days: Available at Primary for fulfillment (arrives the 2nd Saturday)",
]
for c in cross:
    doc.add_paragraph(c, style="List Bullet")
doc.add_paragraph(
    "Plan ahead: anything that needs to be at RMFG TX for Saturday fulfillment "
    "should have been shipped from Woburn 9+ days prior."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Cold Chain & Shipping Basics", level=1)
doc.add_paragraph(
    "Every box ships cold. Understanding the cold chain helps you appreciate "
    "why timing and inventory accuracy matter so much."
)
bullets_cc = [
    "Every order ships with 1x 48oz gel pack baseline; extra gel packs added based on route/weather",
    "Gel packs survive indefinitely below 50°F; above that, heat gain degrades product safety",
    "3-day ground shipping is acceptable in winter for cold states, but NOT in summer or to warm states "
    "(FL, TX, AZ, NV, CA, LA, MS, AL, GA, NM, HI)",
    "We ship from 3 active hubs: Nashville (East Coast), Anaheim (West Coast), Dallas (Central + overflow)",
    "Carriers: OnTrac (regional/cheapest), UPS (national/reliable), FedEx (premium)",
]
for b in bullets_cc:
    doc.add_paragraph(b, style="List Bullet")

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Inventory Management Essentials", level=1)

doc.add_heading("Two Warehouses", level=2)
doc.add_paragraph(
    "Primary (RMFG TX): Fulfillment warehouse — only Primary inventory counts "
    "toward reorder alerts and fulfillment readiness.\n\n"
    "Woburn, MA: Receives bulk cheese, processes accessories, stores and cross-docks "
    "to Texas. Can handle all SKU types (CH-, MT-, AC-)."
)

doc.add_heading("Inventory Columns in Dropbox Snapshot", level=2)
inv_items = [
    "Total = on hand + pending production yield → use for PLANNING",
    "RMFG = actual on hand only → use for RECONCILIATION",
    "Pending = Total − RMFG = expected yield from production not yet completed",
]
for item in inv_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Status Levels", level=2)
t4 = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = t4.rows[0].cells
hdr4[0].text = "Status"
hdr4[1].text = "Meaning"
for row_data in [
    ("OUT OF STOCK", "Zero on hand with active demand — urgent"),
    ("CRITICAL", "≤50% of reorder point or ≤3 days supply"),
    ("REORDER / WARNING", "At or below reorder point, ≤10 days supply"),
    ("OK", "Above reorder point, sufficient stock"),
    ("OVERSTOCK", "More than 3× reorder point — consider reducing orders"),
    ("PLAN", "Will run out within the forecast horizon (1-3 months ahead)"),
]:
    row = t4.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading("Reorder Actions", level=2)
actions = [
    "PO (Purchase Order): Need to order from vendor — meats, finished accessories, cheese wheels",
    "MFG (Manufacturing): Need to cut/wrap/label cheese (from wheels) or process accessories (from bulk)",
    "Transfer: Need to move finished goods from Woburn to RMFG TX",
]
for a in actions:
    doc.add_paragraph(a, style="List Bullet")

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. Paid Bundles & Add-Ons", level=1)
doc.add_paragraph(
    "Some customers purchase add-on bundles (BL- prefix) that expand into multiple "
    "food items. For example, BL-BLR4 expands to 4× CH-BLR + 2× AC-BLBALS. "
    "These items share the same subscription ID as the bundle product."
)
doc.add_paragraph(
    "When checking inventory needs, remember that bundle items are ADDITIONAL "
    "to the standard box contents. A customer with a medium box + BL-BLR4 will "
    "need 7 box items + 6 bundle items = 13 total food items in their shipment."
)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Common Issues & What to Watch For", level=1)

issues = [
    ("Stale items from prior months",
     "When curations rotate, old recipe items sometimes linger on orders. "
     "If you see a cheese that doesn't belong to the current month's recipe, flag it."),
    ("CH-MAFT appearing unexpectedly",
     "CH-MAFT is on the exclusion list. If it shows up outside its native MONG recipe, "
     "it's a data error, not a procurement need."),
    ("Duplicate items on orders",
     "Same food item appearing 2+ times could be: (a) a paid add-on (legitimate — "
     "different subscription IDs), (b) a bundle expansion (legitimate — same sub ID as "
     "BL- product), or (c) a system error (same sub ID, no bundle). Only (c) is a real problem."),
    ("Blank SKUs on promo products",
     "Some promotional Shopify products have blank SKUs. The system determines box size "
     "from the variant title ('Medium (Serves 2-4)' or 'Large (Serves 4-6)')."),
    ("Gift Redemption orders",
     "These cannot be modified automatically. Missing PR-CJAM or tasting guides on "
     "gift orders are a known systemic gap."),
]
for title, desc in issues:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Shortage Resolution Playbook", level=1)
doc.add_paragraph(
    "When a cheese (or any SKU) is short, follow this priority order:"
)
playbook = [
    "Check for incoming product — if a vendor delivery covers the gap, flag as 'incoming' and wait",
    "Use established substitutions — brie family, porter family, alpine/semi-hard (see Section 5)",
    "Reassign PR-CJAM or CEX-EC — move demand off the short cheese onto one with surplus",
    "Cut more wheels — if wheel inventory exists for the short SKU",
    "Recipe change — swap a recipe cheese (last resort, affects all boxes in that curation)",
    "Small custom picks (<10 units) — pull from shelf or substitute per-customer",
]
for i, step in enumerate(playbook, 1):
    doc.add_paragraph(f"{i}. {step}")

# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. Key Contacts & Systems", level=1)
systems = [
    "Shopify — Order management, product catalog, fulfillment tags",
    "Recharge — Subscription management, bundle selections, queued charges",
    "Dropbox — Inventory snapshots (Friday/Monday) from RMFG",
    "RMFG (TX) — Primary fulfillment center, receives cut orders Wednesday",
    "Woburn (MA) — Secondary warehouse for bulk receiving, processing, cross-dock",
    "Inventory Reorder App — Dashboard for demand forecasting, cut orders, reorder alerts",
]
for s in systems:
    doc.add_paragraph(s, style="List Bullet")

# ── Save ────────────────────────────────────────────────────────────
out_dir = os.path.dirname(__file__)
out_path = os.path.join(out_dir, "Elevate_Foods_Procurement_Guide.docx")
doc.save(out_path)
print(f"Document saved: {out_path}")

# ── Upload to Google Drive ──────────────────────────────────────────
try:
    from google_integration import GoogleIntegration
    creds_path = os.path.join(out_dir, "shipping-perfomance-review-accd39ac4b78.json")
    creds_path = os.path.normpath(creds_path)
    print(f"Using credentials: {creds_path}")

    gi = GoogleIntegration(creds_path)
    email = gi.test_connection()
    print(f"Connected as: {email}")

    # Upload as .docx (non-native — no quota issue)
    # Use Google Docs mimetype for convert-on-upload so it opens as a Google Doc
    import io
    from googleapiclient.http import MediaIoBaseUpload

    file_metadata = {
        "name": "Elevate Foods — Procurement & Product Guide (Cheesemonger Onboarding).docx",
        # Upload as non-native .docx — no quota usage for SA
    }
    with open(out_path, "rb") as fh:
        media = MediaIoBaseUpload(
            io.BytesIO(fh.read()),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=True,
        )
        f = gi._drive.files().create(
            body=file_metadata,
            media_body=media,
            fields="id, webViewLink",
        ).execute()

    link = f.get("webViewLink", f"https://docs.google.com/document/d/{f['id']}")
    file_id = f["id"]
    print(f"Uploaded to Google Drive: {link}")

    # Share with kurt@elevatefoods.co
    gi._drive.permissions().create(
        fileId=file_id,
        body={
            "type": "user",
            "role": "writer",
            "emailAddress": "kurt@elevatefoods.co",
        },
        sendNotificationEmail=False,
    ).execute()
    print("Shared with kurt@elevatefoods.co (writer)")

except Exception as e:
    print(f"Drive upload failed: {e}")
    print(f"Document is available locally at: {out_path}")
